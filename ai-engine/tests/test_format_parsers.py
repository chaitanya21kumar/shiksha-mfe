"""Tests for the additional ingestion formats: DOCX, CSV, TXT, Markdown, HTML."""

import docx
from fastapi.testclient import TestClient

from app.ingestion.csv_parser import parse_csv
from app.ingestion.docx_parser import parse_docx
from app.ingestion.html_parser import parse_html
from app.ingestion.md_parser import parse_md
from app.ingestion.txt_parser import parse_txt
from app.main import app

client = TestClient(app)


def _kinds(doc) -> list[str]:
    return [b.kind.value for page in doc.pages for b in page.blocks]


def _block(doc, kind: str):
    return [b for page in doc.pages for b in page.blocks if b.kind.value == kind]


# ---------------------------------------------------------------- TXT
def test_txt_paragraphs_and_lists(tmp_path):
    f = tmp_path / "notes.txt"
    f.write_text("Intro paragraph.\n\nSecond paragraph.\n\n- one\n- two\n- three\n", encoding="utf-8")
    doc = parse_txt(f)

    assert doc.source.format == "txt"
    assert doc.pages[0].kind == "document"
    assert _kinds(doc).count("paragraph") == 2
    assert _block(doc, "list")[0].items == ["one", "two", "three"]


def test_txt_handles_windows_line_endings(tmp_path):
    f = tmp_path / "crlf.txt"
    f.write_bytes(b"First para.\r\n\r\nSecond para.\r\n")
    doc = parse_txt(f)
    assert _kinds(doc).count("paragraph") == 2


# ---------------------------------------------------------------- CSV
def test_csv_becomes_one_table(tmp_path):
    f = tmp_path / "data.csv"
    f.write_text("name,score\nAsha,90\nRavi,85\n", encoding="utf-8")
    doc = parse_csv(f)

    assert doc.source.format == "csv"
    assert doc.pages[0].kind == "sheet"
    table = _block(doc, "table")[0]
    assert table.rows[0] == ["name", "score"]
    assert table.rows[1] == ["Asha", "90"]


# ---------------------------------------------------------------- Markdown
def test_markdown_headings_lists_and_flattened_markup(tmp_path):
    f = tmp_path / "doc.md"
    f.write_text("# Title\n\nA paragraph with **bold** and a [link](http://x).\n\n- a\n- b\n\n## Section\n", encoding="utf-8")
    doc = parse_md(f)

    assert doc.source.format == "md"
    assert doc.source.title == "Title"
    heading = _block(doc, "heading")[0]
    assert heading.text == "Title" and heading.level == 1
    para = _block(doc, "paragraph")[0]
    assert "bold" in para.text and "link" in para.text and "**" not in para.text
    assert _block(doc, "list")[0].items == ["a", "b"]


def test_markdown_table(tmp_path):
    f = tmp_path / "t.md"
    f.write_text("| a | b |\n|---|---|\n| 1 | 2 |\n", encoding="utf-8")
    doc = parse_md(f)

    table = _block(doc, "table")[0]
    assert table.rows[0] == ["a", "b"]
    assert table.rows[1] == ["1", "2"]


# ---------------------------------------------------------------- HTML
def test_html_structure_and_script_stripped(tmp_path):
    f = tmp_path / "page.html"
    f.write_text(
        "<html><head><title>My Page</title></head><body>"
        "<h1>Heading</h1><p>Para text.</p>"
        "<ul><li>one</li><li>two</li></ul>"
        "<table><tr><th>A</th><th>B</th></tr><tr><td>1</td><td>2</td></tr></table>"
        "<script>secretFunction()</script>"
        "</body></html>",
        encoding="utf-8",
    )
    doc = parse_html(f)

    assert doc.source.format == "html"
    assert doc.source.title == "My Page"
    kinds = _kinds(doc)
    assert {"heading", "paragraph", "list", "table"} <= set(kinds)
    assert all("secretFunction" not in (b.text or "") for page in doc.pages for b in page.blocks)


# ---------------------------------------------------------------- DOCX
def _make_docx(path) -> None:
    d = docx.Document()
    d.add_heading("Doc Title", level=1)
    d.add_paragraph("A normal paragraph.")
    d.add_paragraph("First bullet", style="List Bullet")
    d.add_paragraph("Second bullet", style="List Bullet")
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "h1"
    table.rows[0].cells[1].text = "h2"
    table.rows[1].cells[0].text = "v1"
    table.rows[1].cells[1].text = "v2"
    d.save(str(path))


def test_docx_headings_lists_tables(tmp_path):
    f = tmp_path / "doc.docx"
    _make_docx(f)
    doc = parse_docx(f)

    assert doc.source.format == "docx"
    assert doc.pages[0].kind == "document"
    heading = _block(doc, "heading")[0]
    assert heading.text == "Doc Title" and heading.level == 1
    assert _block(doc, "list")[0].items == ["First bullet", "Second bullet"]
    assert _block(doc, "table")[0].rows[0] == ["h1", "h2"]


# ---------------------------------------------------------------- endpoint wiring
def test_ingest_accepts_csv(tmp_path):
    resp = client.post("/ingest", files={"file": ("x.csv", b"a,b\n1,2\n", "text/csv")})
    assert resp.status_code == 200
    assert resp.json()["source"]["format"] == "csv"


def test_ingest_accepts_markdown():
    resp = client.post("/ingest", files={"file": ("n.md", b"# Hi\n\nsome text\n", "text/markdown")})
    assert resp.status_code == 200
    assert resp.json()["source"]["format"] == "md"


def test_ingest_still_rejects_unsupported():
    resp = client.post("/ingest", files={"file": ("a.xyz", b"data", "application/octet-stream")})
    assert resp.status_code == 415
